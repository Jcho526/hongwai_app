#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""分析现有DOC模板的结构"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 加载现有模板
template_file = '评估报告_西医_赵_20251125_214836167_29700_1(1) (1).docx'
doc = Document(template_file)

print("=" * 80)
print("DOC模板结构分析")
print("=" * 80)

# 1. 分析段落
print("\n【段落内容 - 前40个段落】")
for i, para in enumerate(doc.paragraphs[:40]):
    if para.text.strip():
        text_preview = para.text[:60].replace('\n', '\\n')
        print(f"[{i:2d}] 样式: {para.style.name:20s} | {text_preview}")

# 2. 分析表格
print(f"\n【表格信息 - 总共 {len(doc.tables)} 个表格】")
for i, table in enumerate(doc.tables):
    print(f"\n表格 {i}: {len(table.rows)} 行 × {len(table.columns)} 列")
    for j, row in enumerate(table.rows[:5]):  # 显示前5行
        cells_text = [cell.text[:20].replace('\n', ' ') for cell in row.cells]
        print(f"  行 {j}: {cells_text}")
    if len(table.rows) > 5:
        print(f"  ... 还有 {len(table.rows) - 5} 行")

# 3. 分析图片
print("\n【图片信息】")
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        image_count += 1
        print(f"  图片 {image_count}: {rel.target_ref}")
print(f"总图片数: {image_count}")

# 4. 分析样式
print("\n【使用的段落样式】")
styles_used = set()
for para in doc.paragraphs:
    styles_used.add(para.style.name)
for style in sorted(styles_used):
    print(f"  - {style}")

# 5. 分析占位符
print("\n【可能的占位符（包含 {{ 或 [[ 的文本）】")
placeholders = []
for para in doc.paragraphs:
    if '{{' in para.text or '[[' in para.text:
        placeholders.append(para.text)
        print(f"  {para.text[:80]}")

if not placeholders:
    print("  未发现格式化的占位符")

# 6. 分析节信息
print(f"\n【节信息 - 总共 {len(doc.sections)} 个节】")
for i, section in enumerate(doc.sections):
    print(f"节 {i}:")
    print(f"  页面宽度: {section.page_width.cm:.2f} cm")
    print(f"  页面高度: {section.page_height.cm:.2f} cm")
    print(f"  页脚距离: {section.footer_distance.cm:.2f} cm")
    print(f"  页眉距离: {section.header_distance.cm:.2f} cm")

print("\n" + "=" * 80)
print("分析完成！")
