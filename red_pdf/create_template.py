#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建通用的热图分析报告DOCX模板
包含可替换的占位符，方便后续自动化填充
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(paragraph, text, level=1):
    """添加标题样式"""
    paragraph.text = text
    if level == 1:
        paragraph.style = 'Heading 1'
    elif level == 2:
        paragraph.style = 'Heading 2'
    else:
        paragraph.style = 'Heading 3'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_thermal_report_template():
    """创建红外热图分析报告模板"""
    
    doc = Document()
    
    # ==================== 第一页：标题页 ====================
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('红外热像·中医可视化')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    
    # 基本信息
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # 编号
    cell = info_table.rows[0].cells[0]
    cell.text = '编号'
    cell = info_table.rows[0].cells[1]
    cell.text = '{{编号}}'
    
    # 姓名
    cell = info_table.rows[1].cells[0]
    cell.text = '姓名'
    cell = info_table.rows[1].cells[1]
    cell.text = '{{患者名}}'
    
    # 性别和年龄
    cell = info_table.rows[2].cells[0]
    cell.text = '性别'
    cell = info_table.rows[2].cells[1]
    cell.text = '{{性别}}    年龄：{{年龄}}'
    
    # 日期
    cell = info_table.rows[3].cells[0]
    cell.text = '检查日期'
    cell = info_table.rows[3].cells[1]
    cell.text = '{{检查日期}}'
    
    # 添加分页符
    doc.add_page_break()
    
    # ==================== 第二页：目录和注意事项 ====================
    heading = doc.add_paragraph()
    heading.text = '目录'
    heading.style = 'Heading 1'
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('01 / 实时热像中医经络穴位图')
    doc.add_paragraph('01 / 热态分布')
    doc.add_paragraph('02 / 九种体质辨识')
    doc.add_paragraph('03 / 脏腑辨证')
    doc.add_paragraph('03 / 三焦辨证')
    doc.add_paragraph('04 / 经络穴位检测')
    doc.add_paragraph('06 / 气血瘀堵或气血不足')
    doc.add_paragraph('07 / 总体评估')
    doc.add_paragraph('08 / 调理建议')
    
    doc.add_paragraph()  # 空行
    
    # 注意事项
    note_heading = doc.add_paragraph()
    note_heading.text = '重要说明'
    note_heading.style = 'Heading 2'
    
    notes = [
        '饮酒、服用扩血管药物或化疗药物、身体出汗、高热、未能平衡足够时间、扫描姿态不标准，将不同程度影响证素评价分数。',
        '不同环境温度、湿度下进行扫描的结果，会略有差异。',
        '对于手术脏器切除的受检者，相关部位的检查结果及证素评价不作为参考。'
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Number')
    
    doc.add_page_break()
    
    # ==================== 第三页：实时热成像和热态分布 ====================
    heading = doc.add_paragraph()
    heading.text = '实时热成像中医经络穴位图'
    heading.style = 'Heading 1'
    
    doc.add_paragraph('正常人体的温度分布具有一定的稳定性和特征性，机体各部位温度不同，形成了不同的热场。'
                     '当人体某处发生疾病或功能改变时，该处血流量会相应发生变化，导致局部温度改变。'
                     '有的温度升高，有的温度降低。红外热成像技术可以清晰、准确、及时地发现人体组织微小的温度变化。')
    
    doc.add_paragraph('[图片占位符：热图正面图像]')
    doc.add_paragraph('[图片占位符：热图背面图像]')
    
    doc.add_page_break()
    
    # ==================== 热态分布 ====================
    heading = doc.add_paragraph()
    heading.text = '热态分布'
    heading.style = 'Heading 2'
    
    dist_table = doc.add_table(rows=3, cols=2)
    dist_table.style = 'Light Grid Accent 1'
    
    dist_table.rows[0].cells[0].text = '项目'
    dist_table.rows[0].cells[1].text = '数值'
    dist_table.rows[1].cells[0].text = '对称性'
    dist_table.rows[1].cells[1].text = '{{对称性}}'
    dist_table.rows[2].cells[0].text = '规律性'
    dist_table.rows[2].cells[1].text = '{{规律性}}'
    
    doc.add_page_break()
    
    # ==================== 体质辨识 ====================
    heading = doc.add_paragraph()
    heading.text = '热成像中医九种体质辨识模型'
    heading.style = 'Heading 1'
    
    intro = doc.add_paragraph('体质现象是人类生命活动的一种重要表现形式，是指人体生命过程中，'
                             '在先天禀赋和后天获得的基础上所形成的形态结构、生理功能和心理状态方面'
                             '综合的、相对稳定的固有特质。')
    
    doc.add_paragraph('[图片占位符：体质辨识图表]')
    
    # 面部反射区表
    face_heading = doc.add_paragraph()
    face_heading.text = '面部反射区温度分布'
    face_heading.style = 'Heading 3'
    
    face_table = doc.add_table(rows=10, cols=3)
    face_table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = face_table.rows[0].cells
    header_cells[0].text = '部位'
    header_cells[1].text = '温度(℃)'
    header_cells[2].text = '部位'
    
    face_regions = [
        ('首面区', '{{首面区温度}}'),
        ('咽喉区', '{{咽喉区温度}}'),
        ('肺区', '{{肺区温度}}'),
        ('心区', '{{心区温度}}'),
        ('肝区', '{{肝区温度}}'),
        ('脾区', '{{脾区温度}}'),
        ('膀胱', '{{膀胱温度}}'),
        ('小肠区', '{{小肠区温度}}'),
        ('胃区', '{{胃区温度}}'),
    ]
    
    for i, (region, temp) in enumerate(face_regions, 1):
        row = face_table.rows[i]
        row.cells[0].text = region
        row.cells[1].text = temp
    
    doc.add_paragraph('体质类型：{{体质类型}}')
    doc.add_paragraph('体质描述：{{体质描述}}')
    
    doc.add_page_break()
    
    # ==================== 脏腑辨证 ====================
    heading = doc.add_paragraph()
    heading.text = '脏腑辨证'
    heading.style = 'Heading 1'
    
    organ_table = doc.add_table(rows=10, cols=5)
    organ_table.style = 'Light Grid Accent 1'
    
    headers = ['部位', '极大值', '极小值', '平均值', '相对值']
    for i, header in enumerate(headers):
        organ_table.rows[0].cells[i].text = header
    
    organs = [
        ('左胸膺', '{{左胸膺_max}}', '{{左胸膺_min}}', '{{左胸膺_avg}}', '{{左胸膺_rel}}'),
        ('右胸', '{{右胸_max}}', '{{右胸_min}}', '{{右胸_avg}}', '{{右胸_rel}}'),
        ('虚里', '{{虚里_max}}', '{{虚里_min}}', '{{虚里_avg}}', '{{虚里_rel}}'),
        ('胃脘', '{{胃脘_max}}', '{{胃脘_min}}', '{{胃脘_avg}}', '{{胃脘_rel}}'),
        ('左胁', '{{左胁_max}}', '{{左胁_min}}', '{{左胁_avg}}', '{{左胁_rel}}'),
        ('右胁', '{{右胁_max}}', '{{右胁_min}}', '{{右胁_avg}}', '{{右胁_rel}}'),
        ('大腹', '{{大腹_max}}', '{{大腹_min}}', '{{大腹_avg}}', '{{大腹_rel}}'),
        ('小腹', '{{小腹_max}}', '{{小腹_min}}', '{{小腹_avg}}', '{{小腹_rel}}'),
        ('右少腹', '{{右少腹_max}}', '{{右少腹_min}}', '{{右少腹_avg}}', '{{右少腹_rel}}'),
    ]
    
    for i, organ_data in enumerate(organs, 1):
        row = organ_table.rows[i]
        for j, value in enumerate(organ_data):
            row.cells[j].text = value
    
    # 证型描述
    doc.add_paragraph()
    cert_heading = doc.add_paragraph()
    cert_heading.text = '主要证型：{{主要证型}}'
    cert_heading.style = 'Heading 3'
    
    doc.add_paragraph('{{证型详细描述}}')
    
    doc.add_page_break()
    
    # ==================== 三焦分析 ====================
    heading = doc.add_paragraph()
    heading.text = '三焦辨证分析'
    heading.style = 'Heading 1'
    
    tripl_table = doc.add_table(rows=4, cols=5)
    tripl_table.style = 'Light Grid Accent 1'
    
    headers = ['三焦部位', '极大值', '极小值', '平均值', '相对值']
    for i, header in enumerate(headers):
        tripl_table.rows[0].cells[i].text = header
    
    triple_burner = [
        ('上焦', '{{上焦_max}}', '{{上焦_min}}', '{{上焦_avg}}', '{{上焦_rel}}'),
        ('中焦', '{{中焦_max}}', '{{中焦_min}}', '{{中焦_avg}}', '{{中焦_rel}}'),
        ('下焦', '{{下焦_max}}', '{{下焦_min}}', '{{下焦_avg}}', '{{下焦_rel}}'),
    ]
    
    for i, data in enumerate(triple_burner, 1):
        row = tripl_table.rows[i]
        for j, value in enumerate(data):
            row.cells[j].text = value
    
    doc.add_paragraph()
    doc.add_paragraph('三焦分析：{{三焦分析}}')
    doc.add_paragraph('[图片占位符：任脉穴位图]')
    doc.add_paragraph('[图片占位符：督脉穴位图]')
    
    doc.add_page_break()
    
    # ==================== 经络穴位检测 ====================
    heading = doc.add_paragraph()
    heading.text = '经络穴位检测'
    heading.style = 'Heading 1'
    
    meridian_list = [
        '足太阳膀胱经(上)',
        '足太阳膀胱经(下)',
        '足阳明胃经(上)',
        '足太阴脾经(上)',
        '足少阴肾经(上)',
        '手少阴心经',
        '手厥阴心包经',
        '手太阴肺经',
        '手阳明大肠经',
        '手少阳三焦经',
        '手太阳小肠经',
    ]
    
    for meridian in meridian_list:
        para = doc.add_paragraph(style='List Bullet')
        para.text = f'{meridian}：{{{{{{meridian}}_平均温度}}}}'
    
    doc.add_page_break()
    
    # ==================== 气血瘀堵或不足 ====================
    heading = doc.add_paragraph()
    heading.text = '气血瘀堵或气血不足'
    heading.style = 'Heading 1'
    
    intro_text = ('气为血之帅，血为气之母，气能行血，气能摄血。'
                 '血中蕴含的精微物质能濡养肌肉筋脉。'
                 '中医的寒热本质上是各脏腑、组织的温度高低。')
    doc.add_paragraph(intro_text)
    
    doc.add_paragraph('检查日期：{{气血检查日期}}')
    
    # 气血瘀堵区域
    doc.add_paragraph()
    stag_heading = doc.add_paragraph()
    stag_heading.text = '气血瘀堵区域'
    stag_heading.style = 'Heading 3'
    doc.add_paragraph('{{气血瘀堵区域}}')
    
    # 气血不足区域
    doc.add_paragraph()
    insuf_heading = doc.add_paragraph()
    insuf_heading.text = '气血不足区域'
    insuf_heading.style = 'Heading 3'
    doc.add_paragraph('{{气血不足区域}}')
    
    doc.add_paragraph('[图片占位符：手掌热图]')
    doc.add_paragraph('手掌末梢充血性改变，注意血黏稠度。')
    
    doc.add_page_break()
    
    # ==================== 总体评估和调理建议 ====================
    heading = doc.add_paragraph()
    heading.text = '总体评估与调理建议'
    heading.style = 'Heading 1'
    
    # 总体评估
    eval_heading = doc.add_paragraph()
    eval_heading.text = '总体评估'
    eval_heading.style = 'Heading 2'
    
    eval_table = doc.add_table(rows=5, cols=2)
    eval_table.style = 'Light Grid Accent 1'
    
    eval_table.rows[0].cells[0].text = '评估项目'
    eval_table.rows[0].cells[1].text = '评估结果'
    eval_table.rows[1].cells[0].text = '体质类型'
    eval_table.rows[1].cells[1].text = '{{总体体质}}'
    eval_table.rows[2].cells[0].text = '三焦状况'
    eval_table.rows[2].cells[1].text = '{{三焦状况}}'
    eval_table.rows[3].cells[0].text = '气血状况'
    eval_table.rows[3].cells[1].text = '{{气血状况}}'
    eval_table.rows[4].cells[0].text = '主要问题'
    eval_table.rows[4].cells[1].text = '{{主要问题}}'
    
    doc.add_paragraph()
    
    # 调理建议
    suggest_heading = doc.add_paragraph()
    suggest_heading.text = '调理建议'
    suggest_heading.style = 'Heading 2'
    
    # 中医调理
    cm_heading = doc.add_paragraph()
    cm_heading.text = '中医调理原则'
    cm_heading.style = 'Heading 3'
    doc.add_paragraph('{{中医调理原则}}')
    
    # 中药方案
    herbal_heading = doc.add_paragraph()
    herbal_heading.text = '推荐中药方案'
    herbal_heading.style = 'Heading 3'
    doc.add_paragraph('{{推荐中药方案}}')
    
    # 穴位调理
    acupoint_heading = doc.add_paragraph()
    acupoint_heading.text = '穴位调理'
    acupoint_heading.style = 'Heading 3'
    doc.add_paragraph('{{穴位调理}}')
    
    # 药食同源
    food_heading = doc.add_paragraph()
    food_heading.text = '药食同源建议'
    food_heading.style = 'Heading 3'
    doc.add_paragraph('{{药食同源建议}}')
    
    # 起居注意事项
    life_heading = doc.add_paragraph()
    life_heading.text = '起居注意事项'
    life_heading.style = 'Heading 3'
    doc.add_paragraph('{{起居注意事项}}')
    
    doc.add_page_break()
    
    # ==================== 最后一页：免责声明 ====================
    disclaimer_heading = doc.add_paragraph()
    disclaimer_heading.text = '免责声明'
    disclaimer_heading.style = 'Heading 1'
    disclaimer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    disclaimer = ('本报告仅对本次拍摄的红外热像图有效，'
                 '仅表示温度异常所代表的健康风险发生的可能性，'
                 '不能代替临床诊断。'
                 '实际情况以医疗机构诊断为准。')
    doc.add_paragraph(disclaimer)
    
    doc.add_paragraph()
    doc.add_paragraph('评估日期：{{最终评估日期}}')
    doc.add_paragraph('评估医生：{{评估医生}}')
    doc.add_paragraph('医疗机构：{{医疗机构}}')
    
    # 保存文档
    output_file = 'thermal_report_template.docx'
    doc.save(output_file)
    print(f'✅ 通用模板已创建: {output_file}')
    print(f'📋 模板包含以下占位符：')
    print('   - {{编号}}, {{患者名}}, {{性别}}, {{年龄}}, {{检查日期}}')
    print('   - {{对称性}}, {{规律性}}')
    print('   - {{体质类型}}, {{体质描述}}')
    print('   - {{脏腑数据}}, {{三焦数据}}')
    print('   - {{气血瘀堵区域}}, {{气血不足区域}}')
    print('   - {{调理建议}}, {{穴位调理}}, {{药食同源}}')
    print('   - 以及图片占位符：[图片占位符：...]')

if __name__ == '__main__':
    create_thermal_report_template()
