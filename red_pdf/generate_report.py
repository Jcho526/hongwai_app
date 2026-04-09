#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
使用模板生成热图分析报告
演示如何用真实数据填充模板占位符
"""

from docx import Document
from datetime import datetime
import re

class ThermalReportGenerator:
    """热图分析报告生成器"""
    
    def __init__(self, template_path):
        """初始化报告生成器"""
        self.doc = Document(template_path)
        self.placeholders = {}
    
    def set_placeholder(self, key, value):
        """设置占位符的值"""
        self.placeholders[key] = str(value)
    
    def replace_placeholders(self):
        """替换文档中的所有占位符"""
        # 替换段落中的占位符
        for paragraph in self.doc.paragraphs:
            for key, value in self.placeholders.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # 替换表格中的占位符
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in self.placeholders.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
    
    def save(self, output_path):
        """保存生成的报告"""
        self.doc.save(output_path)
        print(f'✅ 报告已生成: {output_path}')

def generate_sample_report():
    """生成示例报告"""
    
    # 创建报告生成器
    generator = ThermalReportGenerator('thermal_report_template.docx')
    
    # 填充基本信息
    generator.set_placeholder('编号', '20260322001')
    generator.set_placeholder('患者名', '赵女士')
    generator.set_placeholder('性别', '女')
    generator.set_placeholder('年龄', '26')
    generator.set_placeholder('检查日期', '2025-11-25')
    
    # 填充热态分布
    generator.set_placeholder('对称性', '左右少腹不对称 0.93')
    generator.set_placeholder('规律性', '正面热序列：肺,左胁,大腹,左胸')
    
    # 填充体质信息
    generator.set_placeholder('体质类型', '阴虚体质，倾向气郁体质')
    generator.set_placeholder('体质描述', '津液不足，阴液损伤，容易出现虚热症状')
    
    # 填充面部反射区温度
    generator.set_placeholder('首面区温度', '29.31')
    generator.set_placeholder('咽喉区温度', '33.68')
    generator.set_placeholder('肺区温度', '34.37')
    generator.set_placeholder('心区温度', '34.24')
    generator.set_placeholder('肝区温度', '34.07')
    generator.set_placeholder('脾区温度', '32.97')
    generator.set_placeholder('膀胱温度', '32.76')
    generator.set_placeholder('小肠区温度', '33.45')
    generator.set_placeholder('胃区温度', '33.22')
    
    # 填充脏腑辨证数据
    organs_data = {
        '左胸膺': (34.78, 23.57, 32.35, 3.20),
        '右胸': (32.10, 27.32, 28.85, -0.30),
        '虚里': (31.73, 27.22, 29.03, -0.12),
        '胃脘': (31.05, 27.13, 28.58, -0.57),
        '左胁': (31.35, 27.57, 29.18, 0.03),
        '右胁': (31.10, 26.95, 28.64, -0.51),
        '大腹': (31.04, 26.75, 29.16, 0.01),
        '小腹': (28.24, 26.81, 27.44, -1.71),
        '右少腹': (28.44, 22.94, 27.17, -1.98),
    }
    
    for organ, (max_t, min_t, avg_t, rel_t) in organs_data.items():
        organ_key = organ.replace('左', '左').replace('右', '右').replace('胁', '胁').replace('膺', '膺').replace('里', '里').replace('脘', '脘').replace('腹', '腹').replace('胸', '胸')
        generator.set_placeholder(f'{organ}_max', f'{max_t:.2f}')
        generator.set_placeholder(f'{organ}_min', f'{min_t:.2f}')
        generator.set_placeholder(f'{organ}_avg', f'{avg_t:.2f}')
        generator.set_placeholder(f'{organ}_rel', f'{rel_t:.2f}')
    
    # 填充三焦数据
    generator.set_placeholder('上焦_max', '33.24')
    generator.set_placeholder('上焦_min', '25.26')
    generator.set_placeholder('上焦_avg', '30.60')
    generator.set_placeholder('上焦_rel', '1.45')
    
    generator.set_placeholder('中焦_max', '31.14')
    generator.set_placeholder('中焦_min', '27.35')
    generator.set_placeholder('中焦_avg', '29.15')
    generator.set_placeholder('中焦_rel', '-0.00')
    
    generator.set_placeholder('下焦_max', '28.89')
    generator.set_placeholder('下焦_min', '25.82')
    generator.set_placeholder('下焦_avg', '27.71')
    generator.set_placeholder('下焦_rel', '-1.44')
    
    # 填充三焦分析
    generator.set_placeholder('三焦分析', '上焦偏热，下焦偏寒。胸部上端代谢热偏高，诊证为湿热阻肺。')
    
    # 填充气血信息
    generator.set_placeholder('气血检查日期', '2025-11-25')
    generator.set_placeholder('气血瘀堵区域', '左后脑[197]，左腹部[134]，左肘[104]，左上臂[147]，左前臂[329]，左后心[94]，右后背区域[452]，右腰[220]')
    generator.set_placeholder('气血不足区域', '右腹部[335]，右上臂[494]，右前臂[343]，右后心[505]，左腰[441]，左后大腿[1612]')
    
    # 填充证型
    generator.set_placeholder('主要证型', '肺阴虚证')
    generator.set_placeholder('证型详细描述', 
        '干咳无痰，或痰少而粘，口燥咽干，形体消瘦，午后潮热，五心烦热，盗汗，颧红。'
        '舌红少津，脉细数。需要滋补肺阴，常用的中药有沙参、麦冬、天冬、玄参、百合等。')
    
    # 填充调理建议
    generator.set_placeholder('中医调理原则', '以养阴润肺，兼顾清热与益气，肺肾同调为原则。')
    
    generator.set_placeholder('推荐中药方案', 
        '1. 百合固金汤：百合、生地黄、熟地黄、麦冬、玄参、当归、白芍、川贝母、桔梗、甘草\n'
        '2. 沙参麦冬汤：沙参、麦冬、玉竹、冬桑叶、生扁豆、天花粉、甘草')
    
    generator.set_placeholder('穴位调理', 
        '1. 肺俞穴：双手拇指顺时针按揉5-10分钟，以酸胀为度。\n'
        '2. 太渊穴：拇指指腹轻柔环形按揉3-5分钟，晨起及睡前各1次。\n'
        '3. 鱼际穴：拇指点压并轻揉2-3分钟，每日3次。')
    
    generator.set_placeholder('药食同源建议',
        '• 肺阴虚燥咳：选麦冬、银耳\n'
        '• 痰多咳喘：用化橘红、杏仁\n'
        '• 脾胃虚寒者：慎用麦冬、生地等寒凉物，可加生姜中和')
    
    generator.set_placeholder('起居注意事项',
        '1. 环境与作息：保湿（室内湿度50%-60%），避免干燥空气刺激；保证22点前入睡\n'
        '2. 运动与呼吸：太极拳、散步每日30分钟；腹式呼吸训练\n'
        '3. 情志管理：悲伤情绪易伤肺，需保持心态平和，可通过音乐、冥想调节')
    
    # 填充总体评估
    generator.set_placeholder('总体体质', '阴虚体质，倾向气郁体质')
    generator.set_placeholder('三焦状况', '上焦热，下焦寒')
    generator.set_placeholder('气血状况', '气血瘀堵与不足并存，末梢循环欠佳')
    generator.set_placeholder('主要问题', '肺阴虚，脾阳虚，小肠实热')
    
    # 填充最终信息
    generator.set_placeholder('最终评估日期', datetime.now().strftime('%Y-%m-%d'))
    generator.set_placeholder('评估医生', '医疗专业人士')
    generator.set_placeholder('医疗机构', '中医诊疗中心')
    
    # 替换所有占位符
    generator.replace_placeholders()
    
    # 生成报告
    output_file = 'thermal_report_sample.docx'
    generator.save(output_file)
    print(f'\n📊 示例报告已生成成功！')
    print(f'📄 文件名: {output_file}')
    print(f'📍 保存位置: c:\\Users\\1\\Desktop\\新\\red_pdf\\{output_file}')

if __name__ == '__main__':
    print('=' * 60)
    print('红外热图分析报告生成器')
    print('=' * 60)
    generate_sample_report()
